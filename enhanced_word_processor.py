#!/usr/bin/env python3
"""
Enhanced Word Controls Processor
Uses the control_mappings.json configuration to properly handle all Word controls.
"""

import os
import json
from datetime import datetime

class EnhancedWordProcessor:
    def __init__(self):
        # Load control mappings from JSON file
        with open('control_mappings.json', 'r') as f:
            self.config = json.load(f)
        self.controls = self.config['controls']
    
    def process_word_template(self, template_path, data, output_path):
        """Process Word template using the control mappings configuration."""
        
        print(f"🔧 Processing Word template with enhanced controls: {template_path}")
        
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Load the template
            doc = Document(template_path)
            
            # Calculate all needed values
            calculations = self.calculate_values(data)
            
            # Build control mappings based on configuration
            control_mappings = self.build_control_mappings(data, calculations)
            
            print("🔄 Processing Word controls with enhanced mappings...")
            
            replacements_made = 0
            
            # Process paragraphs
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                new_text = original_text
                
                for control_name, replacement_value in control_mappings.items():
                    if control_name in new_text:
                        new_text = new_text.replace(control_name, str(replacement_value))
                        replacements_made += 1
                        print(f"   ✅ Replaced '{control_name}' with '{replacement_value}'")
                
                if new_text != original_text:
                    paragraph.clear()
                    paragraph.add_run(new_text)
            
            # Process tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            new_text = original_text
                            
                            for control_name, replacement_value in control_mappings.items():
                                if control_name in new_text:
                                    new_text = new_text.replace(control_name, str(replacement_value))
                                    replacements_made += 1
                                    print(f"   ✅ Replaced '{control_name}' with '{replacement_value}' in table")
                            
                            if new_text != original_text:
                                paragraph.clear()
                                paragraph.add_run(new_text)
            
            # Handle structured content controls if available
            try:
                self.process_content_controls(doc, control_mappings)
            except Exception as e:
                print(f"   ℹ️  Content controls processing: {e}")
            
            print(f"📊 Total replacements made: {replacements_made}")
            
            # Add cost summary tables at the end
            self.add_cost_summary(doc, data)
            
            # Save the processed document
            doc.save(output_path)
            print(f"✅ Enhanced document saved: {output_path}")
            
            return True
            
        except ImportError:
            print("❌ python-docx library not available")
            return False
        except Exception as e:
            print(f"❌ Error processing Word controls: {e}")
            return False
    
    def calculate_values(self, data):
        """Calculate all the values needed for the controls."""
        
        one_time_costs = data.get('oneTimeCosts', [])
        recurring_costs = data.get('recurringCosts', [])
        
        # Basic totals
        one_time_total = sum(item.get('total', 0) for item in one_time_costs)
        recurring_total = sum(item.get('total', 0) for item in recurring_costs)
        
        # Total without VAT (recurring + one time)
        total_excl_vat = one_time_total + recurring_total
        
        # VAT calculation (21%)
        vat_amount = total_excl_vat * 0.21
        
        # Grand total (total + VAT)
        grand_total = total_excl_vat + vat_amount
        
        return {
            'one_time_total': one_time_total,
            'recurring_total': recurring_total,
            'total_excl_vat': total_excl_vat,
            'vat_amount': vat_amount,
            'grand_total': grand_total,
            'current_date': datetime.now().strftime('%d-%m-%Y')
        }
    
    def build_control_mappings(self, data, calculations):
        """Build control mappings based on the configuration."""
        
        mappings = {}
        
        for control_name, config in self.controls.items():
            control_type = config.get('type')
            
            if control_type == 'field':
                # Direct field mapping
                field_name = config.get('value')
                if field_name in data:
                    mappings[control_name] = data[field_name]
                elif field_name == 'companyName':
                    mappings[control_name] = data.get('companyName', '')
                elif field_name == 'contactName':
                    mappings[control_name] = data.get('contactName', '')
                elif field_name == 'address':
                    mappings[control_name] = data.get('address', '')
                elif field_name == 'postalCode':
                    mappings[control_name] = data.get('postalCode', '')
                elif field_name == 'city':
                    mappings[control_name] = data.get('city', '')
                elif field_name == 'companyId':
                    mappings[control_name] = data.get('companyId', '')
                else:
                    mappings[control_name] = ''
                    
            elif control_type == 'calculated':
                # Calculated values
                formula = config.get('formula')
                if formula == 'current_date':
                    mappings[control_name] = calculations['current_date']
                elif formula == 'sum_one_time_costs':
                    mappings[control_name] = f"{calculations['one_time_total']:.2f}"
                elif formula == 'sum_recurring_costs':
                    mappings[control_name] = f"{calculations['recurring_total']:.2f}"
                elif formula == 'recurringandonetimewithoutVAT':
                    mappings[control_name] = f"{calculations['total_excl_vat']:.2f}"
                elif formula == 'VAT':
                    mappings[control_name] = f"{calculations['vat_amount']:.2f}"
                elif formula == 'grandtotal':
                    mappings[control_name] = f"{calculations['grand_total']:.2f}"
                elif formula in ['ammounttimespriceonetimematerial', 'ammounttimespricerecurringmaterial']:
                    # These are per-item calculations - will be handled in item processing
                    mappings[control_name] = ''
                else:
                    mappings[control_name] = ''
                    
            elif control_type == 'list':
                # List processing for items1 and items2
                if config.get('value') == 'oneTimeCosts':
                    mappings[control_name] = self.format_items_table(data.get('oneTimeCosts', []), 'onetime')
                elif config.get('value') == 'recurringCosts':
                    mappings[control_name] = self.format_items_table(data.get('recurringCosts', []), 'recurring')
                else:
                    mappings[control_name] = ''
                    
            elif control_type == 'input':
                # Input fields like description
                if config.get('value') == 'description':
                    mappings[control_name] = data.get('description', 'Quotation generated via web interface')
                else:
                    mappings[control_name] = ''
                    
            else:
                # Unknown or unhandled type
                mappings[control_name] = ''
        
        return mappings
    
    def format_items_table(self, items, item_type):
        """Format items for table display."""
        
        if not items:
            return "Geen items"
        
        formatted_lines = []
        for item in items:
            material = item.get('material', '')
            quantity = item.get('quantity', 0)
            unit_price = item.get('unitPrice', 0)
            total = item.get('total', 0)
            
            # Format as table row or structured text
            line = f"{material} | Aantal: {quantity} | Prijs: €{unit_price:.2f} | Totaal: €{total:.2f}"
            formatted_lines.append(line)
        
        return "\n".join(formatted_lines)
    
    def process_content_controls(self, doc, mappings):
        """Process actual content controls if available."""
        
        # This method tries to access structured document tags
        root = doc.element
        
        # Look for structured document tags (content controls)
        for sdt in root.xpath('.//w:sdt', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
            try:
                # Get control name from alias or tag
                alias_elements = sdt.xpath('.//w:alias/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                tag_elements = sdt.xpath('.//w:tag/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                
                control_name = None
                if alias_elements:
                    control_name = alias_elements[0]
                elif tag_elements:
                    control_name = tag_elements[0]
                
                if control_name and control_name in mappings:
                    # Find text content and replace
                    text_elements = sdt.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                    if text_elements:
                        text_elements[0].text = str(mappings[control_name])
                        print(f"   🎯 Updated content control '{control_name}'")
            
            except Exception:
                continue
    
    def add_cost_summary(self, doc, data):
        """Add cost summary tables at the end."""
        
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            WD_ALIGN_PARAGRAPH = None
        
        one_time_costs = data.get('oneTimeCosts', [])
        recurring_costs = data.get('recurringCosts', [])
        
        if not one_time_costs and not recurring_costs:
            return
        
        print("💰 Adding detailed cost summary tables...")
        
        # Add page break
        doc.add_page_break()
        
        # Add header
        header = doc.add_heading('GEDETAILLEERDE KOSTENSPECIFICATIE', level=1)
        if WD_ALIGN_PARAGRAPH:
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # One-time costs table
        if one_time_costs:
            doc.add_heading('Eenmalige Kosten', level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Module (Product)'
            hdr_cells[1].text = 'Aantal'
            hdr_cells[2].text = 'Prijs per stuk (€)'
            hdr_cells[3].text = 'Totaal (€)'
            
            # Data rows
            for item in one_time_costs:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('material', ''))
                row_cells[1].text = str(item.get('quantity', 0))
                row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                row_cells[3].text = f"€{item.get('total', 0):.2f}"
            
            # Total row
            total = sum(item.get('total', 0) for item in one_time_costs)
            total_row = table.add_row().cells
            total_row[0].text = 'TOTAAL EENMALIG'
            total_row[1].text = ''
            total_row[2].text = ''
            total_row[3].text = f"€{total:.2f}"
            
            doc.add_paragraph('')
        
        # Recurring costs table
        if recurring_costs:
            doc.add_heading('Jaarlijkse Kosten', level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Module (Product)'
            hdr_cells[1].text = 'Aantal'
            hdr_cells[2].text = 'Jaarlijks tarief (€)'
            hdr_cells[3].text = 'Totaal per jaar (€)'
            
            # Data rows
            for item in recurring_costs:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('material', ''))
                row_cells[1].text = str(item.get('quantity', 0))
                row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                row_cells[3].text = f"€{item.get('total', 0):.2f}"
            
            # Total row
            total = sum(item.get('total', 0) for item in recurring_costs)
            total_row = table.add_row().cells
            total_row[0].text = 'TOTAAL JAARLIJKS'
            total_row[1].text = ''
            total_row[2].text = ''
            total_row[3].text = f"€{total:.2f}"

def main():
    """Test the enhanced Word processor."""
    
    processor = EnhancedWordProcessor()
    
    test_data = {
        "companyName": "Test Praktijk BV",
        "contactName": "Dr. Jan Janssen", 
        "address": "Teststraat 123",
        "postalCode": "1234AB",
        "city": "Amsterdam",
        "companyId": "NL123456789B01",
        "description": "Complete IT-oplossing voor tandartspraktijk",
        "oneTimeCosts": [
            {
                "material": "Praktijkbeheersysteem Setup",
                "quantity": 1,
                "unitPrice": 2500.00,
                "total": 2500.00
            },
            {
                "material": "Hardware installatie", 
                "quantity": 3,
                "unitPrice": 400.00,
                "total": 1200.00
            }
        ],
        "recurringCosts": [
            {
                "material": "Software licentie per maand",
                "quantity": 12,
                "unitPrice": 150.00,
                "total": 1800.00
            },
            {
                "material": "Support & onderhoud per jaar",
                "quantity": 1,
                "unitPrice": 600.00,
                "total": 600.00
            }
        ]
    }
    
    success = processor.process_word_template(
        "standaardofferte Compufit NL.docx", 
        test_data, 
        "enhanced_output.docx"
    )
    
    if success:
        print("🎉 Enhanced Word processing test successful!")
        print("📄 Check enhanced_output.docx to see the results")
    else:
        print("❌ Enhanced Word processing test failed")

if __name__ == "__main__":
    main()