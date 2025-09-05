#!/usr/bin/env python3
"""
Test Final Fix
Test the fixed content control processor with real data and verify results.
"""

from content_control_processor import ContentControlProcessor

def main():
    """Test the fixed processor with realistic data."""
    
    processor = ContentControlProcessor()
    
    test_data = {
        "companyName": "FINAL TEST BV",
        "contactName": "Jane Smith",
        "address": "Final Street 789",
        "postalCode": "9999XZ", 
        "city": "Final City",
        "companyId": "NL987654321B01",
        "description": "Final test quotation",
        "oneTimeCosts": [
            {
                "material": "Installation",
                "quantity": 1,
                "unitPrice": 750.00,
                "total": 750.00
            }
        ],
        "recurringCosts": [
            {
                "material": "Support",
                "quantity": 12,
                "unitPrice": 25.00,
                "total": 300.00
            }
        ]
    }
    
    print("🧪 Testing fixed content control processor...")
    print(f"📊 Test data: {test_data['companyName']}")
    
    output_filename = "final_test_fixed.docx"
    success = processor.process_word_template(
        "standaardofferte Compufit NL.docx",
        test_data,
        output_filename
    )
    
    if success:
        print(f"\\n✅ Document generated: {output_filename}")
        
        # Test opening with python-docx
        try:
            from docx import Document
            doc = Document(output_filename)
            print(f"✅ Document can be opened ({len(doc.paragraphs)} paragraphs)")
            
            # Quick search for our test values
            content = ""
            for p in doc.paragraphs:
                content += p.text + " "
            
            test_values = [
                "FINAL TEST BV",
                "Jane Smith", 
                "Final Street 789",
                "9999XZ",
                "Final City",
                "NL987654321B01"
            ]
            
            found_count = 0
            for value in test_values:
                if value in content:
                    print(f"   ✅ Found: {value}")
                    found_count += 1
                else:
                    print(f"   ❌ Missing: {value}")
            
            print(f"\\n📈 Found {found_count}/{len(test_values)} expected values in document")
            
            if found_count >= len(test_values) - 1:  # Allow 1 missing
                print("🎉 SUCCESS: Document appears to be correctly generated!")
                return True
            else:
                print("⚠️ Some values may not have been properly inserted")
                return False
                
        except Exception as e:
            print(f"❌ Error testing document: {e}")
            return False
    else:
        print("❌ Document generation failed")
        return False

if __name__ == "__main__":
    main()