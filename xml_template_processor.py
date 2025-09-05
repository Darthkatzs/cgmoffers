#!/usr/bin/env python3
"""
XML Template Processor
Directly processes the Word document XML to handle broken template tags
and replace them with actual values.
"""

import zipfile
import tempfile
import shutil
import os
from datetime import datetime

class XMLTemplateProcessor:
    def __init__(self):
        pass
    
    def process_word_template(self, template_path, data, output_path):
        """Process the Word template by directly manipulating XML content."""
        
        print(f"🔧 Processing template: {template_path}")
        
        # Prepare replacement data - handle double braces correctly
        replacements = {
            # Handle complete variables with double braces
            '{{praktijknaam}}': data.get('companyName', ''),
            '{{naam}}': data.get('contactName', ''),
            '{{straat}}': data.get('address', ''),
            '{{nummer}}': '',  
            '{{postcode}}': data.get('postalCode', ''),
            '{{stad}}': data.get('city', ''),
            '{{btw}}': data.get('companyId', ''),
            '{{SigB_es_:signer1:signatureblock}}': datetime.now().strftime('%d-%m-%Y'),
        }
        
        # Prepare material lists for Items1 and Items2
        one_time_costs = data.get('oneTimeCosts', [])
        recurring_costs = data.get('recurringCosts', [])
        
        # Format one-time costs list
        items1_text = ""
        if one_time_costs:
            for item in one_time_costs:
                items1_text += f"{item.get('material', '')} - Aantal: {item.get('quantity', 0)} x €{item.get('unitPrice', 0):.2f} = €{item.get('total', 0):.2f}\n"
        
        # Format recurring costs list  
        items2_text = ""
        if recurring_costs:
            for item in recurring_costs:
                items2_text += f"{item.get('material', '')} - Aantal: {item.get('quantity', 0)} x €{item.get('unitPrice', 0):.2f} = €{item.get('total', 0):.2f}\n"
        
        # Add Items1 and Items2 replacements
        replacements['Items1'] = items1_text
        replacements['Items2'] = items2_text
        
        # Also handle broken fragments by reconstructing them
        broken_patterns = {
            ('{{praktijknaam', 'praktijknaam}}'): data.get('companyName', ''),
            ('{{naam', 'naam}}'): data.get('contactName', ''),
            ('{{straat', 'straat}}'): data.get('address', ''),
            ('{{nummer', 'nummer}}'): '',
            ('{{postcode', 'postcode}}'): data.get('postalCode', ''),
            ('{{stad', 'stad}}'): data.get('city', ''),
            ('{{btw', 'btw}}'): data.get('companyId', ''),
            ('{{SigB_es_:signer1:signatureblock', 'SigB_es_:signer1:signatureblock}}'): datetime.now().strftime('%d-%m-%Y'),
        }
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_docx = os.path.join(temp_dir, 'processing.docx')
                shutil.copy2(template_path, temp_docx)
                
                replacements_made = 0
                
                with zipfile.ZipFile(temp_docx, 'r') as input_zip:
                    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as output_zip:
                        
                        for item in input_zip.infolist():
                            data_content = input_zip.read(item.filename)
                            
                            # Process XML files that might contain template variables
                            if item.filename in ['word/document.xml', 'word/header1.xml', 'word/header2.xml', 
                                               'word/header3.xml', 'word/footer1.xml', 'word/footer2.xml', 'word/footer3.xml']:
                                try:
                                    xml_content = data_content.decode('utf-8')
                                    original_content = xml_content
                                    
                                    # First, handle complete variables
                                    for old, new in replacements.items():
                                        if old in xml_content:
                                            xml_content = xml_content.replace(old, str(new))
                                            replacements_made += 1
                                            print(f"   ✅ Replaced {old} with '{new}'")
                                    
                                    # Then, handle broken patterns
                                    for (start_pattern, end_pattern), replacement in broken_patterns.items():
                                        # Look for the pattern and replace both parts
                                        if start_pattern in xml_content and end_pattern in xml_content:
                                            # Replace start pattern
                                            xml_content = xml_content.replace(start_pattern, str(replacement))
                                            # Replace end pattern with empty string
                                            xml_content = xml_content.replace(end_pattern, '')
                                            replacements_made += 1
                                            print(f"   🔨 Fixed broken pattern: {start_pattern}...{end_pattern} -> '{replacement}'")
                                    
                                    # Additional cleanup for any remaining broken fragments and single braces
                                    cleanup_patterns = [
                                        ('praktijk', ''), ('naam}}', ''), ('{{stra', ''), ('raat}}', ''),
                                        ('{{post', ''), ('code}}', ''), ('{{st', ''), ('ad}}', ''),
                                        ('{{bt', ''), ('w}}', ''), ('{{SigB', ''), ('lock}}', ''),
                                        ('{{numm', ''), ('mer}}', ''),
                                        # Remove leftover single braces
                                        ('{praktijknaam}', ''), ('{naam}', ''), ('{straat}', ''), ('{nummer}', ''),
                                        ('{postcode}', ''), ('{stad}', ''), ('{btw}', ''), ('{SigB_es_:signer1:signatureblock}', ''),
                                        # Remove any orphaned curly braces
                                        ('{', ''), ('}', ''),
                                    ]
                                    
                                    for old, new in cleanup_patterns:
                                        if old in xml_content:
                                            xml_content = xml_content.replace(old, new)
                                    
                                    if xml_content != original_content:
                                        print(f"   📝 Modified {item.filename}")
                                    
                                    output_zip.writestr(item, xml_content.encode('utf-8'))
                                    
                                except Exception as e:
                                    print(f"   ⚠️  Error processing {item.filename}: {e}")
                                    output_zip.writestr(item, data_content)
                            else:
                                # Copy other files unchanged
                                output_zip.writestr(item, data_content)
                
                print(f"📊 Total replacements made: {replacements_made}")
                
                # Now add cost information using python-docx
                self.add_cost_tables(output_path, data)
                
                return True
                
        except Exception as e:
            print(f"❌ Error processing template: {e}")
            return False
    
    def add_cost_tables(self, docx_path, data):
        """Add cost tables to the processed document."""
        
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            print("💰 Adding cost tables...")
            
            doc = Document(docx_path)
            
            # Calculate totals
            one_time_costs = data.get('oneTimeCosts', [])
            recurring_costs = data.get('recurringCosts', [])
            one_time_total = sum(item.get('total', 0) for item in one_time_costs)
            recurring_total = sum(item.get('total', 0) for item in recurring_costs)
            
            if one_time_costs or recurring_costs:
                # Add page break
                doc.add_page_break()
                
                # Add costs header
                costs_header = doc.add_heading('KOSTENSPECIFICATIE', level=1)
                costs_header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # One-time costs
                if one_time_costs:
                    doc.add_heading('Eenmalige Kosten', level=2)
                    
                    cost_table = doc.add_table(rows=1, cols=4)
                    cost_table.style = 'Table Grid'
                    
                    # Header row
                    hdr_cells = cost_table.rows[0].cells
                    hdr_cells[0].text = 'Materiaal/Service'
                    hdr_cells[1].text = 'Aantal'
                    hdr_cells[2].text = 'Prijs per stuk'
                    hdr_cells[3].text = 'Totaal'
                    
                    # Add cost rows
                    for item in one_time_costs:
                        row_cells = cost_table.add_row().cells
                        row_cells[0].text = str(item.get('material', ''))
                        row_cells[1].text = str(item.get('quantity', 0))
                        row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                        row_cells[3].text = f"€{item.get('total', 0):.2f}"
                    
                    # Total row
                    total_row = cost_table.add_row().cells
                    total_row[0].text = 'TOTAAL EENMALIG'
                    total_row[1].text = ''
                    total_row[2].text = ''
                    total_row[3].text = f"€{one_time_total:.2f}"
                    
                    doc.add_paragraph('')
                
                # Recurring costs
                if recurring_costs:
                    doc.add_heading('Jaarlijkse Kosten', level=2)
                    
                    recurring_table = doc.add_table(rows=1, cols=4)
                    recurring_table.style = 'Table Grid'
                    
                    # Header row
                    hdr_cells = recurring_table.rows[0].cells
                    hdr_cells[0].text = 'Materiaal/Service'
                    hdr_cells[1].text = 'Aantal'
                    hdr_cells[2].text = 'Prijs per stuk'
                    hdr_cells[3].text = 'Totaal'
                    
                    # Add cost rows
                    for item in recurring_costs:
                        row_cells = recurring_table.add_row().cells
                        row_cells[0].text = str(item.get('material', ''))
                        row_cells[1].text = str(item.get('quantity', 0))
                        row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                        row_cells[3].text = f"€{item.get('total', 0):.2f}"
                    
                    # Total row
                    total_row = recurring_table.add_row().cells
                    total_row[0].text = 'TOTAAL JAARLIJKS'
                    total_row[1].text = ''
                    total_row[2].text = ''
                    total_row[3].text = f"€{recurring_total:.2f}"
                
                # Save the updated document
                doc.save(docx_path)
                print("✅ Cost tables added successfully")
            
        except ImportError:
            print("❌ python-docx not available for adding cost tables")
        except Exception as e:
            print(f"⚠️  Warning: Could not add cost tables: {e}")

def main():
    """Test the XML template processor."""
    
    processor = XMLTemplateProcessor()
    
    test_data = {
        "companyName": "Test Company BV",
        "contactName": "Jan Janssen", 
        "address": "Teststraat 123",
        "postalCode": "1234AB",
        "city": "Amsterdam",
        "companyId": "NL123456789B01",
        "oneTimeCosts": [
            {
                "material": "Setup kosten",
                "quantity": 1,
                "unitPrice": 500.00,
                "total": 500.00
            }
        ],
        "recurringCosts": [
            {
                "material": "Maandelijks onderhoud",
                "quantity": 12,
                "unitPrice": 75.00,
                "total": 900.00
            }
        ]
    }
    
    success = processor.process_word_template(
        "standaardofferte Compufit NL.docx", 
        test_data, 
        "test_output.docx"
    )
    
    if success:
        print("🎉 Template processing test successful!")
        print("📄 Check test_output.docx to see the results")
    else:
        print("❌ Template processing test failed")

if __name__ == "__main__":
    main()