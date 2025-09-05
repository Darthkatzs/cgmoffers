#!/usr/bin/env python3
"""
Test File Opening
Test if the generated files can actually be opened and used.
"""

import os
import zipfile
from docx import Document

def test_file_opening(filename):
    """Test opening a Word document file."""
    
    print(f"🔍 Testing file: {filename}")
    
    if not os.path.exists(filename):
        print(f"❌ File not found: {filename}")
        return False
    
    # Test 1: Check if it's a valid ZIP
    try:
        with zipfile.ZipFile(filename, 'r') as z:
            files = z.namelist()
            print(f"✅ ZIP file valid with {len(files)} files")
    except Exception as e:
        print(f"❌ ZIP error: {e}")
        return False
    
    # Test 2: Try opening with python-docx
    try:
        doc = Document(filename)
        print(f"✅ python-docx can open file")
        print(f"   📄 {len(doc.paragraphs)} paragraphs")
        print(f"   📊 {len(doc.tables)} tables")
        
        # Test reading some content
        content_found = False
        for paragraph in doc.paragraphs[:10]:
            if paragraph.text.strip():
                print(f"   📝 Sample text: '{paragraph.text[:50]}...'")
                content_found = True
                break
        
        if not content_found:
            print("   ⚠️  No text content found in first 10 paragraphs")
            
        return True
        
    except Exception as e:
        print(f"❌ python-docx error: {e}")
        return False

def main():
    """Test the most recent generated files."""
    
    test_files = [
        'content_control_output.docx',  # Most recent
        'minimal_test.docx',  # If it exists
        'standaardofferte Compufit NL.docx'  # Original for comparison
    ]
    
    for filename in test_files:
        if os.path.exists(filename):
            print("\n" + "="*50)
            success = test_file_opening(filename)
            if success:
                print("✅ PASSED")
            else:
                print("❌ FAILED")
        else:
            print(f"⚠️  File not found: {filename}")

if __name__ == "__main__":
    main()