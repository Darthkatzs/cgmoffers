#!/usr/bin/env python3
"""
Content Control Processor
Properly processes Word content controls (Structured Document Tags) using XML manipulation.
"""

import zipfile
import xml.etree.ElementTree as ET
import tempfile
import shutil
import os
import json
import copy
from datetime import datetime

class ContentControlProcessor:
    def __init__(self):
        # Load control mappings from JSON file
        with open('control_mappings.json', 'r') as f:
            self.config = json.load(f)
        self.controls = self.config['controls']
        # Feature flag: enable repeating sections with robust cloning
        self.enable_repeating_sections = True
    
    def process_word_template(self, template_path, data, output_path):
        """Process Word template by directly manipulating content controls in XML."""
        
        print(f"🔧 Processing Word template content controls: {template_path}")
        
        # Calculate all needed values
        calculations = self.calculate_values(data)
        
        # Build control mappings
        control_mappings = self.build_control_mappings(data, calculations)
        
        try:
            with tempfile.TemporaryDirectory() as temp_dir:
                temp_docx = os.path.join(temp_dir, 'processing.docx')
                shutil.copy2(template_path, temp_docx)
                
                replacements_made = 0
                
                with zipfile.ZipFile(temp_docx, 'r') as input_zip:
                    with zipfile.ZipFile(output_path, 'w', zipfile.ZIP_DEFLATED) as output_zip:
                        
                        for item in input_zip.infolist():
                            data_content = input_zip.read(item.filename)
                            
                            # Process XML files that might contain content controls
                            if item.filename in ['word/document.xml', 'word/header1.xml', 'word/header2.xml', 
                                               'word/header3.xml', 'word/footer1.xml', 'word/footer2.xml', 'word/footer3.xml']:
                                try:
                                    xml_content = data_content.decode('utf-8')
                                    
                                    # Process content controls in this XML
                                    modified_xml, changes = self.process_content_controls_xml(xml_content, control_mappings, data)
                                    
                                    if changes > 0:
                                        print(f"   📄 {item.filename}: {changes} controls updated")
                                        replacements_made += changes
                                        output_zip.writestr(item, modified_xml.encode('utf-8'))
                                    else:
                                        output_zip.writestr(item, data_content)
                                    
                                except Exception as e:
                                    print(f"   ⚠️  Error processing {item.filename}: {e}")
                                    output_zip.writestr(item, data_content)
                            else:
                                # Copy other files unchanged
                                output_zip.writestr(item, data_content)
                
                print(f"📊 Total content controls updated: {replacements_made}")
                
                # Add detailed cost summary using python-docx
                self.add_cost_summary_to_docx(output_path, data)
                
                return True
                
        except Exception as e:
            print(f"❌ Error processing content controls: {e}")
            return False
    
    def process_content_controls_xml(self, xml_content, control_mappings, data):
        """Process content controls in XML content with repeating section support."""
        
        changes_made = 0
        
        # First handle repeating sections (disabled via feature flag while we analyze)
        if self.enable_repeating_sections:
            xml_content, repeating_changes = self.process_repeating_sections(xml_content, data)
            changes_made += repeating_changes
        
        try:
            # Parse XML with namespace handling
            # Register all original namespaces to preserve formatting
            ET.register_namespace('wpc', 'http://schemas.microsoft.com/office/word/2010/wordprocessingCanvas')
            ET.register_namespace('cx', 'http://schemas.microsoft.com/office/drawing/2014/chartex')
            ET.register_namespace('cx1', 'http://schemas.microsoft.com/office/drawing/2015/9/8/chartex')
            ET.register_namespace('cx2', 'http://schemas.microsoft.com/office/drawing/2015/10/21/chartex')
            ET.register_namespace('cx3', 'http://schemas.microsoft.com/office/drawing/2016/5/9/chartex')
            ET.register_namespace('cx4', 'http://schemas.microsoft.com/office/drawing/2016/5/10/chartex')
            ET.register_namespace('cx5', 'http://schemas.microsoft.com/office/drawing/2016/5/11/chartex')
            ET.register_namespace('cx6', 'http://schemas.microsoft.com/office/drawing/2016/5/12/chartex')
            ET.register_namespace('cx7', 'http://schemas.microsoft.com/office/drawing/2016/5/13/chartex')
            ET.register_namespace('cx8', 'http://schemas.microsoft.com/office/drawing/2016/5/14/chartex')
            ET.register_namespace('mc', 'http://schemas.openxmlformats.org/markup-compatibility/2006')
            ET.register_namespace('aink', 'http://schemas.microsoft.com/office/drawing/2016/ink')
            ET.register_namespace('am3d', 'http://schemas.microsoft.com/office/drawing/2017/model3d')
            ET.register_namespace('o', 'urn:schemas-microsoft-com:office:office')
            ET.register_namespace('oel', 'http://schemas.microsoft.com/office/2019/extlst')
            ET.register_namespace('r', 'http://schemas.openxmlformats.org/officeDocument/2006/relationships')
            ET.register_namespace('m', 'http://schemas.openxmlformats.org/officeDocument/2006/math')
            ET.register_namespace('v', 'urn:schemas-microsoft-com:vml')
            ET.register_namespace('wp14', 'http://schemas.microsoft.com/office/word/2010/wordprocessingDrawing')
            ET.register_namespace('wp', 'http://schemas.openxmlformats.org/drawingml/2006/wordprocessingDrawing')
            ET.register_namespace('w10', 'urn:schemas-microsoft-com:office:word')
            ET.register_namespace('w', 'http://schemas.openxmlformats.org/wordprocessingml/2006/main')
            ET.register_namespace('w14', 'http://schemas.microsoft.com/office/word/2010/wordml')
            ET.register_namespace('w15', 'http://schemas.microsoft.com/office/word/2012/wordml')
            ET.register_namespace('w16cex', 'http://schemas.microsoft.com/office/word/2018/wordml/cex')
            ET.register_namespace('w16cid', 'http://schemas.microsoft.com/office/word/2016/wordml/cid')
            ET.register_namespace('w16', 'http://schemas.microsoft.com/office/word/2018/wordml')
            ET.register_namespace('w16du', 'http://schemas.microsoft.com/office/word/2023/wordml/word16du')
            ET.register_namespace('w16sdtdh', 'http://schemas.microsoft.com/office/word/2020/wordml/sdtdatahash')
            ET.register_namespace('w16sdtfl', 'http://schemas.microsoft.com/office/word/2024/wordml/sdtformatlock')
            ET.register_namespace('w16se', 'http://schemas.microsoft.com/office/word/2015/wordml/symex')
            ET.register_namespace('wpg', 'http://schemas.microsoft.com/office/word/2010/wordprocessingGroup')
            ET.register_namespace('wpi', 'http://schemas.microsoft.com/office/word/2010/wordprocessingInk')
            ET.register_namespace('wne', 'http://schemas.microsoft.com/office/word/2006/wordml')
            ET.register_namespace('wps', 'http://schemas.microsoft.com/office/word/2010/wordprocessingShape')
            
            root = ET.fromstring(xml_content)
            # Build parent map to detect ancestry
            parent_map = {c: p for p in root.iter() for c in p}
            
            # Find all Structured Document Tags (content controls)
            w_ns = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            
            # Track instances of duplicate control names
            control_instances = {}
            
            row_fields = {
                'Module', 'Aantal', 'éénmalige setupkost', 'calctotaalsetup', 'Jaarlijks', 'calctotaaljaarlijks'
            }

            def is_inside_repeating_section(node):
                w15_ns = '{http://schemas.microsoft.com/office/word/2012/wordml}'
                cur = node
                while cur is not None:
                    if cur.tag == f'{w_ns}sdt':
                        pr = cur.find(f'{w_ns}sdtPr')
                        if pr is not None and pr.find(f'{w15_ns}repeatingSectionItem') is not None:
                            return True
                    cur = parent_map.get(cur)
                return False

            for sdt in root.iter(f'{w_ns}sdt'):
                try:
                    # Find the SDT properties to get the control name
                    sdt_pr = sdt.find(f'{w_ns}sdtPr')
                    if sdt_pr is not None:
                        
                        # Try to find alias first, then tag
                        control_name = None
                        
                        alias_elem = sdt_pr.find(f'{w_ns}alias')
                        if alias_elem is not None:
                            control_name = alias_elem.get(f'{w_ns}val')
                        
                        if not control_name:
                            tag_elem = sdt_pr.find(f'{w_ns}tag')
                            if tag_elem is not None:
                                control_name = tag_elem.get(f'{w_ns}val')
                        
                        # If we found a control name
                        if control_name:
                            # Always avoid overriding row controls inside repeating sections
                            inside_repeating = is_inside_repeating_section(sdt)
                            if inside_repeating and control_name in row_fields:
                                print(f"      ⏭️  Skipping control '{control_name}' inside repeating section")
                                continue
                            # Track which instance this is
                            if control_name not in control_instances:
                                control_instances[control_name] = 0
                            control_instances[control_name] += 1
                            instance_num = control_instances[control_name]
                            
                            # Get the replacement value (context-aware for table fields)
                            replacement_value = self.get_contextual_value(control_name, instance_num, control_mappings, data)
                            
                            if replacement_value is not None:
                                # Find the content part of the SDT and update it
                                sdt_content = sdt.find(f'{w_ns}sdtContent')
                                if sdt_content is not None:
                                    # Clear all existing text content and replace with new value
                                    text_elements = []
                                    for t in sdt_content.iter(f'{w_ns}t'):
                                        text_elements.append(t)
                                    
                                    if text_elements:
                                        # Clear all text elements
                                        for t in text_elements:
                                            t.text = ""
                                        
                                        # Set the first one to our replacement value
                                        text_elements[0].text = str(replacement_value)
                                        changes_made += 1
                                        print(f"      ✅ Updated control '{control_name}' (instance {instance_num}) -> '{replacement_value}'")
                                    else:
                                        # No text elements found, create new ones
                                        new_p = ET.SubElement(sdt_content, f'{w_ns}p')
                                        new_r = ET.SubElement(new_p, f'{w_ns}r')
                                        new_t = ET.SubElement(new_r, f'{w_ns}t')
                                        new_t.text = str(replacement_value)
                                        changes_made += 1
                                        print(f"      ✅ Created new content for '{control_name}' (instance {instance_num}) -> '{replacement_value}'")
                
                except Exception as e:
                    print(f"      ⚠️  Error processing SDT: {e}")
                    continue
            
            # Convert back to string while preserving the original XML declaration
            modified_xml = ET.tostring(root, encoding='unicode')
            
            # Preserve the original XML declaration
            if xml_content.startswith('<?xml'):
                xml_declaration = xml_content.split('?>')[0] + '?>'
                # Remove any XML declaration from the generated content
                if modified_xml.startswith('<?xml'):
                    modified_xml = modified_xml.split('?>', 1)[1]
                modified_xml = xml_declaration + modified_xml
            
            return modified_xml, changes_made
            
        except ET.ParseError as e:
            print(f"   ❌ XML Parse Error: {e}")
            return xml_content, 0
        except Exception as e:
            print(f"   ❌ Error: {e}")
            return xml_content, 0
    
    def get_contextual_value(self, control_name, instance_num, control_mappings, data):
        """Get contextual value for a control based on its instance number and context."""
        
        # Handle table fields that need context-aware values
        if control_name == 'Module':
            if instance_num == 1:
                # First instance: use first one-time cost
                one_time_costs = data.get('oneTimeCosts', [])
                if one_time_costs and len(one_time_costs) >= 1:
                    return one_time_costs[0].get('material', '')
            elif instance_num == 2:
                # Second instance: use first recurring cost
                recurring_costs = data.get('recurringCosts', [])
                if recurring_costs and len(recurring_costs) >= 1:
                    return recurring_costs[0].get('material', '')
            return ''
            
        elif control_name == 'Aantal':
            if instance_num == 1:
                # First instance: use first one-time cost
                one_time_costs = data.get('oneTimeCosts', [])
                if one_time_costs and len(one_time_costs) >= 1:
                    return str(one_time_costs[0].get('quantity', ''))
            elif instance_num == 2:
                # Second instance: use first recurring cost
                recurring_costs = data.get('recurringCosts', [])
                if recurring_costs and len(recurring_costs) >= 1:
                    return str(recurring_costs[0].get('quantity', ''))
            return ''
            
        # Handle price fields
        elif control_name == 'éénmalige setupkost':
            if instance_num == 1:
                # First instance: use first one-time cost
                one_time_costs = data.get('oneTimeCosts', [])
                if one_time_costs and len(one_time_costs) >= 1:
                    return f"€{one_time_costs[0].get('unitPrice', 0):.2f}"
            return '€0.00'
            
        elif control_name == 'calctotaalsetup':
            if instance_num == 1:
                # First instance: use first one-time cost
                one_time_costs = data.get('oneTimeCosts', [])
                if one_time_costs and len(one_time_costs) >= 1:
                    total = one_time_costs[0].get('quantity', 0) * one_time_costs[0].get('unitPrice', 0)
                    return f"€{total:.2f}"
            return '€0.00'
            
        elif control_name == 'Jaarlijks':
            if instance_num == 2:
                # Second instance: use first recurring cost
                recurring_costs = data.get('recurringCosts', [])
                if recurring_costs and len(recurring_costs) >= 1:
                    return f"€{recurring_costs[0].get('unitPrice', 0):.2f}"
            return '€0.00'
            
        elif control_name == 'calctotaaljaarlijks':
            if instance_num == 2:
                # Second instance: use first recurring cost
                recurring_costs = data.get('recurringCosts', [])
                if recurring_costs and len(recurring_costs) >= 1:
                    total = recurring_costs[0].get('quantity', 0) * recurring_costs[0].get('unitPrice', 0)
                    return f"€{total:.2f}"
            return '€0.00'
        
        # For non-contextual controls, use the standard mapping
        elif control_name in control_mappings:
            return control_mappings[control_name]
            
        # No mapping found
        return None
    
    def process_repeating_sections(self, xml_content, data):
        """Handle Word repeating section content controls for items1 and items2."""
        
        changes_made = 0
        
        try:
            # Handle items1 (one-time costs)
            one_time_costs = data.get('oneTimeCosts', [])
            print(f"   🔍 Processing items1 with {len(one_time_costs)} items: {one_time_costs}")
            if one_time_costs:
                xml_content, items1_changes = self.duplicate_repeating_section(
                    xml_content, 'items1', one_time_costs
                )
                changes_made += items1_changes
                print(f"   ✅ items1 processing: {items1_changes} changes made")
            
            # Handle items2 (recurring costs)
            recurring_costs = data.get('recurringCosts', [])
            print(f"   🔍 Processing items2 with {len(recurring_costs)} items: {recurring_costs}")
            if recurring_costs:
                xml_content, items2_changes = self.duplicate_repeating_section(
                    xml_content, 'items2', recurring_costs
                )
                changes_made += items2_changes
                print(f"   ✅ items2 processing: {items2_changes} changes made")
                
        except Exception as e:
            print(f"   ⚠️  Error processing repeating sections: {e}")
            import traceback
            traceback.print_exc()
        
        return xml_content, changes_made
    
    def duplicate_repeating_section(self, xml_content, section_name, items):
        """Duplicate a repeating section for each item by cloning the entire repeatingSectionItem SDT per row."""
        changes_made = 0
        try:
            print(f"   🔍 Looking for section: {section_name}")

            # Preserve original XML declaration if present
            xml_decl = None
            if xml_content.startswith('<?xml'):
                xml_decl = xml_content.split('?>', 1)[0] + '?>'

            root = ET.fromstring(xml_content)
            w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
            w15 = '{http://schemas.microsoft.com/office/word/2012/wordml}'

            # Parent map for removals
            parent_map = {c: p for p in root.iter() for c in p}

            # 1) Find the section container SDT by its tag value
            section_sdt = None
            for sdt in root.iter(f'{w}sdt'):
                pr = sdt.find(f'{w}sdtPr')
                if pr is None:
                    continue
                tag = pr.find(f'{w}tag')
                if tag is not None and tag.get(f'{w}val') == section_name:
                    section_sdt = sdt
                    break

            if section_sdt is None:
                print(f"   ❌ No section container found for {section_name}")
                return xml_content, 0

            print(f"   ✅ Found section container for {section_name}")

            section_content = section_sdt.find(f'{w}sdtContent')
            if section_content is None:
                print(f"   ❌ Section {section_name} has no sdtContent")
                return xml_content, 0

            # 2) Find the repeatingSectionItem SDT(s) inside the section
            item_sdts = []
            for nested in section_content.iter(f'{w}sdt'):
                npr = nested.find(f'{w}sdtPr')
                if npr is None:
                    continue
                if any(True for _ in npr.iter(f'{w15}repeatingSectionItem')):
                    item_sdts.append(nested)

            if not item_sdts:
                print(f"   ❌ No repeatingSectionItem found in {section_name}")
                return xml_content, 0

            # Use the first item SDT as the template
            template_item_sdt = item_sdts[0]
            template_parent = parent_map.get(template_item_sdt)
            if template_parent is None:
                print(f"   ❌ Could not locate parent for repeating item in {section_name}")
                return xml_content, 0

            print(f"   ✅ Found repeatingSectionItem template for {section_name}")

            # Remove all existing repeating item SDTs under the parent
            for node in list(template_parent):
                # Remove siblings that are repeating items
                if node.tag == f'{w}sdt':
                    pr = node.find(f'{w}sdtPr')
                    if pr is not None and any(True for _ in pr.iter(f'{w15}repeatingSectionItem')):
                        template_parent.remove(node)

            # Create a clone for each item and set row controls
            for idx, item in enumerate(items, start=1):
                print(f"   🔄 Processing item {idx}: {item}")
                clone = copy.deepcopy(template_item_sdt)

                content = clone.find(f'{w}sdtContent')
                if content is None:
                    continue

                # Set values within this clone's subtree
                material = str(item.get('material', ''))
                qty = str(item.get('quantity', ''))
                unit = float(item.get('unitPrice', 0))
                total = float(item.get('quantity', 0)) * float(item.get('unitPrice', 0))

                self._set_control_value_in_element(content, 'Module', material)
                self._set_control_value_in_element(content, 'Aantal', qty)
                if section_name == 'items1':
                    self._set_control_value_in_element(content, 'éénmalige setupkost', f"€{unit:.2f}")
                    self._set_control_value_in_element(content, 'calctotaalsetup', f"€{total:.2f}")
                else:
                    self._set_control_value_in_element(content, 'Jaarlijks', f"€{unit:.2f}")
                    self._set_control_value_in_element(content, 'calctotaaljaarlijks', f"€{total:.2f}")

                template_parent.append(clone)
                changes_made += 1

            # Convert back to string, preserving XML declaration if present
            modified = ET.tostring(root, encoding='unicode')
            if xml_decl and not modified.startswith('<?xml'):
                modified = xml_decl + modified

            print(f"   ✅ Replaced section {section_name} with {len(items)} items")
            return modified, changes_made

        except Exception as e:
            print(f"   ⚠️  Error duplicating section {section_name}: {e}")
            import traceback
            traceback.print_exc()
            return xml_content, 0

    def _set_control_value_in_element(self, element, control_name, value):
        """Set the text of a content control (by alias or tag) within the given element subtree."""
        w = '{http://schemas.openxmlformats.org/wordprocessingml/2006/main}'
        updated = False
        for sdt in element.iter(f'{w}sdt'):
            try:
                pr = sdt.find(f'{w}sdtPr')
                if pr is None:
                    continue
                name = None
                alias = pr.find(f'{w}alias')
                if alias is not None:
                    name = alias.get(f'{w}val')
                if not name:
                    tag = pr.find(f'{w}tag')
                    if tag is not None:
                        name = tag.get(f'{w}val')
                if name != control_name:
                    continue
                content = sdt.find(f'{w}sdtContent')
                if content is None:
                    continue
                # Find first text run and set it
                t_elem = None
                for t in content.iter(f'{w}t'):
                    t_elem = t
                    break
                if t_elem is None:
                    # Create a new paragraph/run/text if needed
                    p = ET.SubElement(content, f'{w}p')
                    r = ET.SubElement(p, f'{w}r')
                    t_elem = ET.SubElement(r, f'{w}t')
                t_elem.text = str(value)
                updated = True
            except Exception:
                continue
        return updated
    
    def replace_control_in_xml(self, xml_content, control_name, value):
        """Replace a specific content control in XML content."""
        
        import re
        
        print(f"       🔍 Looking for control: {control_name} with value: {value}")
        
        # Pattern to find content control by alias or tag
        patterns = [
            rf'(<w:sdt[^>]*>.*?<w:alias w:val="{re.escape(control_name)}"[^>]*/>.*?<w:sdtContent>.*?<w:t[^>]*>)[^<]*(</w:t>.*?</w:sdtContent>.*?</w:sdt>)',
            rf'(<w:sdt[^>]*>.*?<w:tag w:val="{re.escape(control_name)}"[^>]*/>.*?<w:sdtContent>.*?<w:t[^>]*>)[^<]*(</w:t>.*?</w:sdtContent>.*?</w:sdt>)'
        ]
        
        found = False
        for i, pattern in enumerate(patterns):
            match = re.search(pattern, xml_content, re.DOTALL)
            if match:
                print(f"       ✅ Found control {control_name} using pattern {i+1}")
                replacement = match.group(1) + str(value) + match.group(2)
                xml_content = xml_content.replace(match.group(0), replacement)
                found = True
                break
        
        if not found:
            print(f"       ❌ Control {control_name} not found in XML")
        
        return xml_content
    
    def calculate_values(self, data):
        """Calculate all the values needed for the controls."""
        
        one_time_costs = data.get('oneTimeCosts', [])
        recurring_costs = data.get('recurringCosts', [])
        
        # Basic totals
        one_time_total = sum(item.get('total', 0) for item in one_time_costs)
        recurring_total = sum(item.get('total', 0) for item in recurring_costs)
        
        # Total without VAT (recurring + one time)
        total_excl_vat = one_time_total + recurring_total
        
        # VAT calculation (21%)
        vat_amount = total_excl_vat * 0.21
        
        # Grand total (total + VAT)
        grand_total = total_excl_vat + vat_amount
        
        return {
            'one_time_total': one_time_total,
            'recurring_total': recurring_total,
            'total_excl_vat': total_excl_vat,
            'vat_amount': vat_amount,
            'grand_total': grand_total,
            'current_date': datetime.now().strftime('%d-%m-%Y')
        }
    
    def build_control_mappings(self, data, calculations):
        """Build control mappings based on the configuration."""
        
        mappings = {}
        
        for control_name, config in self.controls.items():
            control_type = config.get('type')
            
            if control_type == 'field':
                # Direct field mapping
                field_name = config.get('value')
                if field_name == 'companyName':
                    mappings[control_name] = data.get('companyName', '')
                elif field_name == 'contactName':
                    mappings[control_name] = data.get('contactName', '')
                elif field_name == 'address':
                    mappings[control_name] = data.get('address', '')
                elif field_name == 'postalCode':
                    mappings[control_name] = data.get('postalCode', '')
                elif field_name == 'city':
                    mappings[control_name] = data.get('city', '')
                elif field_name == 'companyId':
                    mappings[control_name] = data.get('companyId', '')
                # Handle special table fields that pull from cost items
                elif field_name == 'modulenname':
                    # Get module name from first one-time cost item
                    one_time_costs = data.get('oneTimeCosts', [])
                    if one_time_costs:
                        mappings[control_name] = one_time_costs[0].get('material', '')
                    else:
                        mappings[control_name] = ''
                elif field_name == 'itemammount':
                    # Get quantity from first one-time cost item  
                    one_time_costs = data.get('oneTimeCosts', [])
                    if one_time_costs:
                        mappings[control_name] = str(one_time_costs[0].get('quantity', ''))
                    else:
                        mappings[control_name] = ''
                elif field_name == 'annualmaterialcost':
                    # Get unit price from first recurring cost item
                    recurring_costs = data.get('recurringCosts', [])
                    if recurring_costs:
                        mappings[control_name] = f"€{recurring_costs[0].get('unitPrice', 0):.2f}"
                    else:
                        mappings[control_name] = ''
                else:
                    mappings[control_name] = data.get(field_name, '')
                    
            elif control_type == 'calculated':
                # Calculated values
                formula = config.get('formula')
                value = config.get('value')  # Some calculated fields use 'value' instead of 'formula'
                
                if formula == 'current_date':
                    mappings[control_name] = calculations['current_date']
                elif formula == 'sum_one_time_costs':
                    mappings[control_name] = f"{calculations['one_time_total']:.2f}"
                elif formula == 'sum_recurring_costs':
                    mappings[control_name] = f"{calculations['recurring_total']:.2f}"
                elif formula == 'recurringandonetimewithoutVAT':
                    mappings[control_name] = f"{calculations['total_excl_vat']:.2f}"
                elif formula == 'VAT':
                    mappings[control_name] = f"{calculations['vat_amount']:.2f}"
                elif formula == 'grandtotal':
                    mappings[control_name] = f"{calculations['grand_total']:.2f}"
                # Table-specific calculated fields
                elif formula == 'ammounttimespriceonetimematerial':
                    # Calculate total for first one-time cost item
                    one_time_costs = data.get('oneTimeCosts', [])
                    if one_time_costs:
                        item = one_time_costs[0]
                        total = item.get('quantity', 0) * item.get('unitPrice', 0)
                        mappings[control_name] = f"€{total:.2f}"
                    else:
                        mappings[control_name] = '€0.00'
                elif formula == 'ammounttimespricerecurringmaterial':
                    # Calculate total for first recurring cost item
                    recurring_costs = data.get('recurringCosts', [])
                    if recurring_costs:
                        item = recurring_costs[0]
                        total = item.get('quantity', 0) * item.get('unitPrice', 0)
                        mappings[control_name] = f"€{total:.2f}"
                    else:
                        mappings[control_name] = '€0.00'
                # Handle calculated fields that use 'value' instead of 'formula'
                elif value == 'totalsetup':
                    # Unit price for first one-time cost item
                    one_time_costs = data.get('oneTimeCosts', [])
                    if one_time_costs:
                        unit_price = one_time_costs[0].get('unitPrice', 0)
                        mappings[control_name] = f"€{unit_price:.2f}"
                    else:
                        mappings[control_name] = '€0.00'
                else:
                    mappings[control_name] = ''
                    
            elif control_type == 'list':
                # List processing for items1 and items2
                if config.get('value') == 'oneTimeCosts':
                    mappings[control_name] = self.format_items_list(data.get('oneTimeCosts', []))
                elif config.get('value') == 'recurringCosts':
                    mappings[control_name] = self.format_items_list(data.get('recurringCosts', []))
                else:
                    mappings[control_name] = ''
                    
            elif control_type == 'input':
                # Input fields like description
                if config.get('value') == 'description':
                    mappings[control_name] = data.get('description', '')
                else:
                    mappings[control_name] = ''
                    
            else:
                # Unknown or unhandled type
                mappings[control_name] = ''
        
        return mappings
    
    def format_items_list(self, items):
        """Format items for list display."""
        
        if not items:
            return "Geen items"
        
        formatted_lines = []
        for item in items:
            material = item.get('material', '')
            quantity = item.get('quantity', 0)
            unit_price = item.get('unitPrice', 0)
            total = item.get('total', 0)
            
            line = f"• {material} - Aantal: {quantity} x €{unit_price:.2f} = €{total:.2f}"
            formatted_lines.append(line)
        
        return "\n".join(formatted_lines)
    
    def add_cost_summary_to_docx(self, docx_path, data):
        """Add cost summary tables using python-docx."""
        
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            print("💰 Adding cost summary tables...")
            
            doc = Document(docx_path)
            
            one_time_costs = data.get('oneTimeCosts', [])
            recurring_costs = data.get('recurringCosts', [])
            
            if one_time_costs or recurring_costs:
                # Add page break
                doc.add_page_break()
                
                # Add header
                header = doc.add_heading('KOSTEN SPECIFICATIE', level=1)
                if WD_ALIGN_PARAGRAPH:
                    header.alignment = WD_ALIGN_PARAGRAPH.CENTER
                
                # Add tables (same as before)
                if one_time_costs:
                    doc.add_heading('Eenmalige Kosten Detail', level=2)
                    
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Module'
                    hdr_cells[1].text = 'Aantal'
                    hdr_cells[2].text = 'Prijs per stuk'
                    hdr_cells[3].text = 'Totaal'
                    
                    for item in one_time_costs:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(item.get('material', ''))
                        row_cells[1].text = str(item.get('quantity', 0))
                        row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                        row_cells[3].text = f"€{item.get('total', 0):.2f}"
                    
                    total = sum(item.get('total', 0) for item in one_time_costs)
                    total_row = table.add_row().cells
                    total_row[0].text = 'TOTAAL EENMALIG'
                    total_row[1].text = ''
                    total_row[2].text = ''
                    total_row[3].text = f"€{total:.2f}"
                
                if recurring_costs:
                    doc.add_heading('Jaarlijkse Kosten Detail', level=2)
                    
                    table = doc.add_table(rows=1, cols=4)
                    table.style = 'Table Grid'
                    
                    hdr_cells = table.rows[0].cells
                    hdr_cells[0].text = 'Module'
                    hdr_cells[1].text = 'Aantal'
                    hdr_cells[2].text = 'Jaarlijks'
                    hdr_cells[3].text = 'Totaal'
                    
                    for item in recurring_costs:
                        row_cells = table.add_row().cells
                        row_cells[0].text = str(item.get('material', ''))
                        row_cells[1].text = str(item.get('quantity', 0))
                        row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                        row_cells[3].text = f"€{item.get('total', 0):.2f}"
                    
                    total = sum(item.get('total', 0) for item in recurring_costs)
                    total_row = table.add_row().cells
                    total_row[0].text = 'TOTAAL JAARLIJKS'
                    total_row[1].text = ''
                    total_row[2].text = ''
                    total_row[3].text = f"€{total:.2f}"
                
                # Save the document
                doc.save(docx_path)
                print("✅ Cost summary tables added")
                
        except ImportError:
            print("⚠️  python-docx not available for cost summary")
        except Exception as e:
            print(f"⚠️  Error adding cost summary: {e}")

def main():
    """Test the content control processor."""
    
    processor = ContentControlProcessor()
    
    test_data = {
        "companyName": "TEST PRAKTIJK BV",
        "contactName": "Dr. Test Person", 
        "address": "Test Street 123",
        "postalCode": "1234AB",
        "city": "Test City",
        "companyId": "TEST123456",
        "description": "Test quotation description",
        "oneTimeCosts": [
            {
                "material": "Setup",
                "quantity": 1,
                "unitPrice": 1000.00,
                "total": 1000.00
            }
        ],
        "recurringCosts": [
            {
                "material": "Monthly service",
                "quantity": 12,
                "unitPrice": 100.00,
                "total": 1200.00
            }
        ]
    }
    
    success = processor.process_word_template(
        "standaardofferte Compufit NL.docx", 
        test_data, 
        "content_control_output.docx"
    )
    
    if success:
        print("🎉 Content control processing test successful!")
        print("📄 Check content_control_output.docx to see the results")
    else:
        print("❌ Content control processing test failed")

if __name__ == "__main__":
    main()