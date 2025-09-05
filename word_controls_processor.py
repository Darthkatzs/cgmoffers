#!/usr/bin/env python3
"""
Word Controls Processor
Works with Word form controls/content controls instead of template variables.
This is much cleaner and more reliable than XML parsing.
"""

import os
import shutil
from datetime import datetime

class WordControlsProcessor:
    def __init__(self):
        pass
    
    def process_word_template(self, template_path, data, output_path):
        """Process Word template using python-docx to work with controls."""
        
        print(f"🔧 Processing Word template with controls: {template_path}")
        
        try:
            from docx import Document
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            
            # Load the template
            doc = Document(template_path)
            
            # Calculate totals for cost calculations
            one_time_costs = data.get('oneTimeCosts', [])
            recurring_costs = data.get('recurringCosts', [])
            one_time_total = sum(item.get('total', 0) for item in one_time_costs)
            recurring_total = sum(item.get('total', 0) for item in recurring_costs)
            total_excl_vat = one_time_total + recurring_total
            vat_amount = total_excl_vat * 0.21  # 21% VAT
            grand_total = total_excl_vat + vat_amount
            
            # Prepare the data mapping - using all controls found in your document
            control_mappings = {
                # Basic company information (using both alias and tag names)
                'praktijk': data.get('companyName', ''),
                'companyName': data.get('companyName', ''),
                'praktijknaam': data.get('companyName', ''),
                'naam': data.get('contactName', ''),
                'contactName': data.get('contactName', ''),
                'adres': data.get('address', ''),
                'address': data.get('address', ''),
                'straat': data.get('address', ''),
                'nummer': '',  # House number if needed
                'postcode': data.get('postalCode', ''),
                'postalCode': data.get('postalCode', ''),
                'stad': data.get('city', ''),
                'city': data.get('city', ''),
                'btw': data.get('companyId', ''),
                'companyId': data.get('companyId', ''),
                
                # Date fields
                'date': datetime.now().strftime('%d-%m-%Y'),
                'SigB_es_:signer1:signatureblock': datetime.now().strftime('%d-%m-%Y'),
                'signer1': datetime.now().strftime('%d-%m-%Y'),
                'SigB_es': datetime.now().strftime('%d-%m-%Y'),
                'signatureblock': datetime.now().strftime('%d-%m-%Y'),
                
                # Cost list controls
                'items1': self.format_cost_list(data.get('oneTimeCosts', [])),
                'Items1': self.format_cost_list(data.get('oneTimeCosts', [])),
                'items2': self.format_cost_list(data.get('recurringCosts', [])),
                'Items2': self.format_cost_list(data.get('recurringCosts', [])),
                
                # Total calculations
                'totaaleenmalig': f"{one_time_total:.2f}",
                'totaaljaarlijks': f"{recurring_total:.2f}",
                'calctotaalsetup': f"{one_time_total:.2f}",
                'calctotaaljaarlijks': f"{recurring_total:.2f}",
                'total': f"{total_excl_vat:.2f}",
                'vat': f"{vat_amount:.2f}",
                'grandtotal': f"{grand_total:.2f}",
                
                # Description field
                'beschrijving': 'Quotation generated via web interface',
                
                # Module and pricing fields (these might be used in tables)
                'Module': '',
                'Aantal': '',
                'éénmalige setupkost': '',
                'Jaarlijks': '',
                
                # Company/practice name variations
                'Bedrijf': data.get('companyName', ''),
                'Naam': data.get('contactName', ''),
                'Praktijknaam': data.get('companyName', ''),
            }
            
            print("🔄 Processing Word controls...")
            
            # Method 1: Try to find and replace content controls
            replacements_made = 0
            
            # Process paragraphs - look for control content
            for paragraph in doc.paragraphs:
                original_text = paragraph.text
                new_text = original_text
                
                for control_name, replacement_value in control_mappings.items():
                    # Look for the control name in the paragraph
                    if control_name in new_text:
                        new_text = new_text.replace(control_name, str(replacement_value))
                        replacements_made += 1
                        print(f"   ✅ Replaced '{control_name}' with '{replacement_value}'")
                
                # If text changed, update the paragraph
                if new_text != original_text:
                    paragraph.clear()
                    paragraph.add_run(new_text)
            
            # Process tables - look for control content
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for paragraph in cell.paragraphs:
                            original_text = paragraph.text
                            new_text = original_text
                            
                            for control_name, replacement_value in control_mappings.items():
                                if control_name in new_text:
                                    new_text = new_text.replace(control_name, str(replacement_value))
                                    replacements_made += 1
                                    print(f"   ✅ Replaced '{control_name}' with '{replacement_value}' in table")
                            
                            if new_text != original_text:
                                paragraph.clear()
                                paragraph.add_run(new_text)
            
            # Method 2: Try to access actual content controls (if they exist)
            try:
                # This works if the document has proper content controls
                from docx.oxml.ns import qn
                from docx.oxml import parse_xml
                
                # Look for structured document tags (content controls)
                root = doc.element
                for sdt in root.xpath('.//w:sdt', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'}):
                    # Try to find the control name and replace content
                    try:
                        # Get the alias/tag name
                        alias_elements = sdt.xpath('.//w:alias/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        tag_elements = sdt.xpath('.//w:tag/@w:val', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                        
                        control_name = None
                        if alias_elements:
                            control_name = alias_elements[0]
                        elif tag_elements:
                            control_name = tag_elements[0]
                        
                        if control_name and control_name in control_mappings:
                            # Find the text content and replace it
                            text_elements = sdt.xpath('.//w:t', namespaces={'w': 'http://schemas.openxmlformats.org/wordprocessingml/2006/main'})
                            if text_elements:
                                text_elements[0].text = str(control_mappings[control_name])
                                replacements_made += 1
                                print(f"   🎯 Updated control '{control_name}' with '{control_mappings[control_name]}'")
                    
                    except Exception as e:
                        # Skip this control if there's an issue
                        continue
                        
            except Exception as e:
                print(f"   ℹ️  Content controls method not available: {e}")
            
            print(f"📊 Total replacements made: {replacements_made}")
            
            # Add cost summary tables at the end
            self.add_cost_summary(doc, data)
            
            # Save the processed document
            doc.save(output_path)
            print(f"✅ Processed document saved: {output_path}")
            
            return True
            
        except ImportError:
            print("❌ python-docx library not available")
            return False
        except Exception as e:
            print(f"❌ Error processing Word controls: {e}")
            return False
    
    def format_cost_list(self, costs):
        """Format a list of costs for display."""
        if not costs:
            return "Geen items"
        
        formatted_list = []
        for item in costs:
            line = f"{item.get('material', '')} - Aantal: {item.get('quantity', 0)} x €{item.get('unitPrice', 0):.2f} = €{item.get('total', 0):.2f}"
            formatted_list.append(line)
        
        return "\n".join(formatted_list)
    
    def add_cost_summary(self, doc, data):
        """Add cost summary tables to the document."""
        
        try:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
        except ImportError:
            WD_ALIGN_PARAGRAPH = None
        
        one_time_costs = data.get('oneTimeCosts', [])
        recurring_costs = data.get('recurringCosts', [])
        
        if not one_time_costs and not recurring_costs:
            return
        
        print("💰 Adding cost summary tables...")
        
        # Add page break
        doc.add_page_break()
        
        # Add header
        header = doc.add_heading('KOSTENDETAILS', level=1)
        if WD_ALIGN_PARAGRAPH:
            header.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # One-time costs table
        if one_time_costs:
            doc.add_heading('Eenmalige Kosten', level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Materiaal/Service'
            hdr_cells[1].text = 'Aantal'
            hdr_cells[2].text = 'Prijs per stuk'
            hdr_cells[3].text = 'Totaal'
            
            # Data rows
            for item in one_time_costs:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('material', ''))
                row_cells[1].text = str(item.get('quantity', 0))
                row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                row_cells[3].text = f"€{item.get('total', 0):.2f}"
            
            # Total row
            total = sum(item.get('total', 0) for item in one_time_costs)
            total_row = table.add_row().cells
            total_row[0].text = 'TOTAAL EENMALIG'
            total_row[1].text = ''
            total_row[2].text = ''
            total_row[3].text = f"€{total:.2f}"
            
            doc.add_paragraph('')
        
        # Recurring costs table
        if recurring_costs:
            doc.add_heading('Jaarlijkse Kosten', level=2)
            
            table = doc.add_table(rows=1, cols=4)
            table.style = 'Table Grid'
            
            # Header row
            hdr_cells = table.rows[0].cells
            hdr_cells[0].text = 'Materiaal/Service'
            hdr_cells[1].text = 'Aantal'
            hdr_cells[2].text = 'Prijs per stuk'
            hdr_cells[3].text = 'Totaal'
            
            # Data rows
            for item in recurring_costs:
                row_cells = table.add_row().cells
                row_cells[0].text = str(item.get('material', ''))
                row_cells[1].text = str(item.get('quantity', 0))
                row_cells[2].text = f"€{item.get('unitPrice', 0):.2f}"
                row_cells[3].text = f"€{item.get('total', 0):.2f}"
            
            # Total row
            total = sum(item.get('total', 0) for item in recurring_costs)
            total_row = table.add_row().cells
            total_row[0].text = 'TOTAAL JAARLIJKS'
            total_row[1].text = ''
            total_row[2].text = ''
            total_row[3].text = f"€{total:.2f}"

def main():
    """Test the Word controls processor."""
    
    processor = WordControlsProcessor()
    
    test_data = {
        "companyName": "Test Company BV",
        "contactName": "Jan Janssen", 
        "address": "Teststraat 123",
        "postalCode": "1234AB",
        "city": "Amsterdam",
        "companyId": "NL123456789B01",
        "oneTimeCosts": [
            {
                "material": "Setup kosten",
                "quantity": 1,
                "unitPrice": 500.00,
                "total": 500.00
            }
        ],
        "recurringCosts": [
            {
                "material": "Maandelijks onderhoud",
                "quantity": 12,
                "unitPrice": 75.00,
                "total": 900.00
            }
        ]
    }
    
    success = processor.process_word_template(
        "standaardofferte Compufit NL.docx", 
        test_data, 
        "test_controls_output.docx"
    )
    
    if success:
        print("🎉 Word controls processing test successful!")
        print("📄 Check test_controls_output.docx to see the results")
    else:
        print("❌ Word controls processing test failed")

if __name__ == "__main__":
    main()