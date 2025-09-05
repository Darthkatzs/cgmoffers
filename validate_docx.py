#!/usr/bin/env python3
"""
Validate DOCX Structure
Check if generated Word documents have valid XML structure and can be opened.
"""

import zipfile
import xml.etree.ElementTree as ET
import os

def validate_docx(filename):
    """Validate a DOCX file structure."""
    
    print(f"🔍 Validating DOCX file: {filename}")
    
    if not os.path.exists(filename):
        print(f"❌ File not found: {filename}")
        return False
    
    print(f"📏 File size: {os.path.getsize(filename)} bytes")
    
    try:
        # Test if it can be opened as a ZIP file
        with zipfile.ZipFile(filename, 'r') as docx_zip:
            file_list = docx_zip.namelist()
            print(f"📁 ZIP contains {len(file_list)} files")
            
            # Check for essential files
            essential_files = [
                'word/document.xml',
                '[Content_Types].xml',
                '_rels/.rels'
            ]
            
            missing_files = []
            for essential_file in essential_files:
                if essential_file not in file_list:
                    missing_files.append(essential_file)
            
            if missing_files:
                print(f"❌ Missing essential files: {missing_files}")
                return False
            else:
                print("✅ All essential files present")
            
            # Check main document XML
            try:
                document_xml = docx_zip.read('word/document.xml').decode('utf-8')
                print(f"📄 Document XML size: {len(document_xml)} characters")
                
                # Try to parse the XML
                ET.fromstring(document_xml)
                print("✅ Document XML is valid")
                
                # Check if it starts properly
                if document_xml.startswith('<?xml'):
                    print("✅ XML has proper header")
                else:
                    print("⚠️  XML missing header")
                    
            except ET.ParseError as e:
                print(f"❌ Document XML parse error: {e}")
                return False
            except Exception as e:
                print(f"❌ Document XML read error: {e}")
                return False
            
            # Check Content Types
            try:
                content_types = docx_zip.read('[Content_Types].xml').decode('utf-8')
                ET.fromstring(content_types)
                print("✅ Content Types XML is valid")
            except Exception as e:
                print(f"❌ Content Types XML error: {e}")
                return False
                
            print("✅ DOCX file appears to be valid")
            return True
            
    except zipfile.BadZipFile:
        print("❌ File is not a valid ZIP archive")
        return False
    except Exception as e:
        print(f"❌ Validation error: {e}")
        return False

def test_with_python_docx(filename):
    """Test opening with python-docx library."""
    
    print(f"\n📚 Testing with python-docx: {filename}")
    
    try:
        from docx import Document
        doc = Document(filename)
        
        paragraph_count = len(doc.paragraphs)
        table_count = len(doc.tables)
        
        print(f"✅ Opened with python-docx successfully")
        print(f"📄 Contains {paragraph_count} paragraphs and {table_count} tables")
        return True
        
    except Exception as e:
        print(f"❌ python-docx error: {e}")
        return False

def main():
    """Test validation on recent generated files."""
    
    test_files = [
        'content_control_output.docx',
        'Offerte_Test_Company_BV_20250905_215004.docx',
        'Offerte_bedrijfsnaamX_20250905_215042.docx'
    ]
    
    for test_file in test_files:
        if os.path.exists(test_file):
            print("\n" + "="*60)
            print(f"TESTING: {test_file}")
            print("="*60)
            
            is_valid = validate_docx(test_file)
            if is_valid:
                test_with_python_docx(test_file)
        else:
            print(f"⚠️  File not found: {test_file}")

if __name__ == "__main__":
    main()