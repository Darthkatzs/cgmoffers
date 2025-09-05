#!/usr/bin/env python3
"""
Create a clean Word template from scratch that works with docxtemplater.
This bypasses all the broken tag issues by creating a minimal template.
"""

from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
import sys
import os

def create_clean_template(filename="clean_template.docx"):
    """Create a clean Word template with proper variables."""
    
    print(f"🔧 Creating clean template: {filename}")
    print("=" * 60)
    
    # Create a new document
    doc = Document()
    
    # Add title
    title = doc.add_heading('OFFERTE', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # Add a page break
    doc.add_page_break()
    
    # Company Information Section
    doc.add_heading('Klantgegevens', level=1)
    
    # Create a simple table for company info
    table = doc.add_table(rows=6, cols=2)
    table.style = 'Table Grid'
    
    # Fill the table
    cells = [
        ('Bedrijf:', '{companyName}'),
        ('Contactpersoon:', '{contactName}'),
        ('Adres:', '{address}'),
        ('Postcode:', '{postalCode}'),
        ('Plaats:', '{city}'),
        ('Bedrijfs-ID:', '{companyId}')
    ]
    
    for i, (label, value) in enumerate(cells):
        table.cell(i, 0).text = label
        table.cell(i, 1).text = value
    
    # Add date
    doc.add_paragraph(f'Datum: {{date}}')
    
    doc.add_paragraph('')  # Space
    
    # One-time costs section
    doc.add_heading('Eenmalige Kosten', level=2)
    doc.add_paragraph('{{#hasOneTimeCosts}}')
    doc.add_paragraph('{{#oneTimeCosts}}')
    doc.add_paragraph('• {{material}} - Aantal: {{quantity}} x €{{unitPrice}} = €{{total}}')
    doc.add_paragraph('{{/oneTimeCosts}}')
    doc.add_paragraph('Totaal Eenmalig: €{{oneTimeTotal}}')
    doc.add_paragraph('{{/hasOneTimeCosts}}')
    
    doc.add_paragraph('')  # Space
    
    # Recurring costs section
    doc.add_heading('Jaarlijkse Kosten', level=2) 
    doc.add_paragraph('{{#hasRecurringCosts}}')
    doc.add_paragraph('{{#recurringCosts}}')
    doc.add_paragraph('• {{material}} - Aantal: {{quantity}} x €{{unitPrice}} = €{{total}}')
    doc.add_paragraph('{{/recurringCosts}}')
    doc.add_paragraph('Totaal Jaarlijks: €{{recurringTotal}}')
    doc.add_paragraph('{{/hasRecurringCosts}}')
    
    doc.add_paragraph('')  # Space
    
    # Footer information
    doc.add_heading('Met vriendelijke groet,', level=2)
    doc.add_paragraph('')
    doc.add_paragraph('Compufit')
    doc.add_paragraph('Datum: {date}')
    
    # Save the document
    doc.save(filename)
    
    print(f"✅ Clean template created: {filename}")
    print("🎯 This template should work without duplicate tag errors!")
    
    return True

def main():
    """Main function."""
    output_file = "clean_template.docx"
    
    if len(sys.argv) > 1:
        output_file = sys.argv[1]
    
    success = create_clean_template(output_file)
    
    if success:
        print(f"\n🎉 Success! Use this template instead:")
        print(f"1. Rename '{output_file}' to 'standaardofferte Compufit NL.docx'")
        print(f"2. Or update your script to use '{output_file}'")
        print(f"3. Test the PDF generation")
        print(f"\n💡 This template has clean variables that won't cause duplicate tag errors.")
    
    return 0

if __name__ == "__main__":
    try:
        exit(main())
    except ImportError:
        print("❌ Error: python-docx library not found!")
        print("Install it with: pip install python-docx")
        exit(1)