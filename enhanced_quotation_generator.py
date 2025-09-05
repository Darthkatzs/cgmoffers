#!/usr/bin/env python3
"""
Enhanced Quotation Generator
Web server that processes quotations and generates PDFs from Word templates,
with robust handling of broken template tags.
"""

import http.server
import socketserver
import json
import os
import sys
import tempfile
import subprocess
from urllib.parse import parse_qs, urlparse
from robust_template_processor import RobustTemplateProcessor

class QuotationHTTPRequestHandler(http.server.SimpleHTTPRequestHandler):
    def __init__(self, *args, **kwargs):
        self.template_processor = RobustTemplateProcessor()
        super().__init__(*args, **kwargs)
    
    def end_headers(self):
        """Add CORS headers for all responses."""
        self.send_header('Access-Control-Allow-Origin', '*')
        self.send_header('Access-Control-Allow-Methods', 'GET, POST, OPTIONS')
        self.send_header('Access-Control-Allow-Headers', 'Content-Type')
        super().end_headers()
    
    def do_OPTIONS(self):
        """Handle preflight requests."""
        self.send_response(200)
        self.end_headers()
    
    def do_POST(self):
        """Handle POST requests for quotation generation."""
        
        if self.path == '/generate-quotation':
            try:
                # Read the request data
                content_length = int(self.headers['Content-Length'])
                post_data = self.rfile.read(content_length)
                quotation_data = json.loads(post_data.decode('utf-8'))
                
                print(f"🎯 Generating quotation for: {quotation_data.get('companyName', 'Unknown')}")
                
                # Process the quotation
                result = self.generate_quotation_pdf(quotation_data)
                
                if result['success']:
                    self.send_response(200)
                    self.send_header('Content-Type', 'application/json')
                    self.end_headers()
                    
                    response = {
                        'success': True,
                        'message': 'Quotation generated successfully',
                        'filename': result['filename']
                    }
                    self.wfile.write(json.dumps(response).encode())
                else:
                    self.send_error(500, result['error'])
            
            except Exception as e:
                print(f"❌ Error generating quotation: {e}")
                self.send_error(500, f"Error generating quotation: {str(e)}")
        else:
            self.send_error(404, "Endpoint not found")
    
    def generate_quotation_pdf(self, data):
        """Generate a PDF quotation from the Word template."""
        
        template_file = "standaardofferte Compufit NL.docx"
        
        if not os.path.exists(template_file):
            return {'success': False, 'error': f'Template file {template_file} not found'}
        
        try:
            # Create a temporary processed template
            with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as temp_template:
                temp_template_path = temp_template.name
            
            # Process the template to fix any broken tags
            print("🔧 Processing template for broken tags...")
            success = self.template_processor.process_docx_template(template_file, temp_template_path)
            
            if not success:
                return {'success': False, 'error': 'Failed to process template'}
            
            # Now use Python to fill the template (since JavaScript docxtemplater has issues)
            filled_doc_path = self.fill_template_with_python(temp_template_path, data)
            
            if not filled_doc_path:
                return {'success': False, 'error': 'Failed to fill template'}
            
            # Convert to PDF (if needed)
            pdf_path = self.convert_to_pdf(filled_doc_path, data.get('companyName', 'quotation'))
            
            # Clean up temporary files
            try:
                os.unlink(temp_template_path)
                if filled_doc_path != pdf_path:
                    os.unlink(filled_doc_path)
            except:
                pass
            
            return {'success': True, 'filename': os.path.basename(pdf_path)}
            
        except Exception as e:
            return {'success': False, 'error': str(e)}
    
    def fill_template_with_python(self, template_path, data):
        """Fill the template using Python docxtemplater alternative."""
        
        try:
            from docx import Document
            from datetime import datetime
            
            # Open the processed template
            doc = Document(template_path)
            
            # Prepare the replacement data
            replacements = {
                '{companyName}': data.get('companyName', ''),
                '{contactName}': data.get('contactName', ''),
                '{address}': data.get('address', ''),
                '{postalCode}': data.get('postalCode', ''),
                '{city}': data.get('city', ''),
                '{companyId}': data.get('companyId', ''),
                '{date}': datetime.now().strftime('%d-%m-%Y')
            }
            
            # Process one-time costs
            one_time_costs = data.get('oneTimeCosts', [])
            one_time_total = sum(item.get('total', 0) for item in one_time_costs)
            
            # Process recurring costs
            recurring_costs = data.get('recurringCosts', [])
            recurring_total = sum(item.get('total', 0) for item in recurring_costs)
            
            # Add cost totals
            replacements['{oneTimeTotal}'] = f"{one_time_total:.2f}"
            replacements['{recurringTotal}'] = f"{recurring_total:.2f}"
            
            # Replace text in all paragraphs
            for paragraph in doc.paragraphs:
                for key, value in replacements.items():
                    if key in paragraph.text:
                        paragraph.text = paragraph.text.replace(key, str(value))
            
            # Replace text in tables
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for key, value in replacements.items():
                            if key in cell.text:
                                cell.text = cell.text.replace(key, str(value))
            
            # Handle cost lists (simple approach)
            cost_text = ""
            if one_time_costs:
                cost_text += "EENMALIGE KOSTEN:\n"
                for item in one_time_costs:
                    cost_text += f"• {item.get('material', '')} - Aantal: {item.get('quantity', 0)} x €{item.get('unitPrice', 0):.2f} = €{item.get('total', 0):.2f}\n"
                cost_text += f"Totaal Eenmalig: €{one_time_total:.2f}\n\n"
            
            if recurring_costs:
                cost_text += "JAARLIJKSE KOSTEN:\n"
                for item in recurring_costs:
                    cost_text += f"• {item.get('material', '')} - Aantal: {item.get('quantity', 0)} x €{item.get('unitPrice', 0):.2f} = €{item.get('total', 0):.2f}\n"
                cost_text += f"Totaal Jaarlijks: €{recurring_total:.2f}\n"
            
            # Add cost details to the document (simple approach - add at end)
            if cost_text:
                doc.add_paragraph(cost_text)
            
            # Save filled document
            output_path = f"quotation_{data.get('companyName', 'unknown').replace(' ', '_')}_{datetime.now().strftime('%Y%m%d_%H%M%S')}.docx"
            doc.save(output_path)
            
            print(f"✅ Template filled: {output_path}")
            return output_path
            
        except ImportError:
            print("❌ python-docx not installed. Install with: pip install python-docx")
            return None
        except Exception as e:
            print(f"❌ Error filling template: {e}")
            return None
    
    def convert_to_pdf(self, docx_path, company_name):
        """Convert Word document to PDF (if tools are available)."""
        
        # For now, just return the Word document
        # PDF conversion can be added later with tools like:
        # - LibreOffice: libreoffice --headless --convert-to pdf file.docx
        # - pandoc: pandoc file.docx -o file.pdf
        
        print(f"📄 Document ready: {docx_path}")
        return docx_path

def main():
    """Start the enhanced quotation server."""
    
    port = 8000
    if len(sys.argv) > 1:
        port = int(sys.argv[1])
    
    print(f"🚀 Starting Enhanced Quotation Generator Server")
    print(f"📍 Port: {port}")
    print(f"🌐 URL: http://localhost:{port}")
    print(f"📋 Template: standaardofferte Compufit NL.docx")
    print("=" * 60)
    
    try:
        with socketserver.TCPServer(("", port), QuotationHTTPRequestHandler) as httpd:
            print(f"✅ Server running! Access your quotation system at:")
            print(f"   http://localhost:{port}")
            print(f"\n💡 The server will automatically process your existing Word template")
            print(f"   to handle any broken template tags.")
            print(f"\n🛑 Press Ctrl+C to stop the server")
            
            httpd.serve_forever()
            
    except KeyboardInterrupt:
        print(f"\n👋 Server stopped")
    except Exception as e:
        print(f"❌ Server error: {e}")

if __name__ == "__main__":
    main()