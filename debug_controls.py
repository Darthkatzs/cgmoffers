#!/usr/bin/env python3
"""
Debug Word Controls
Debug what's actually happening with the control replacements.
"""

from datetime import datetime

def debug_word_template():
    """Debug the Word template processing."""
    
    try:
        from docx import Document
        
        print("🔍 DEBUG: Loading Word template...")
        doc = Document("standaardofferte Compufit NL.docx")
        
        # Test data
        test_data = {
            "companyName": "DEBUG_COMPANY",
            "contactName": "DEBUG_CONTACT", 
            "address": "DEBUG_ADDRESS",
            "postalCode": "DEBUG_POSTAL",
            "city": "DEBUG_CITY",
            "companyId": "DEBUG_BTW",
            "description": "DEBUG_DESCRIPTION"
        }
        
        # Simple control mappings for debugging
        control_mappings = {
            'praktijk': 'DEBUG_COMPANY',
            'companyName': 'DEBUG_COMPANY',
            'praktijknaam': 'DEBUG_COMPANY',
            'naam': 'DEBUG_CONTACT',
            'contactName': 'DEBUG_CONTACT',
            'adres': 'DEBUG_ADDRESS',
            'address': 'DEBUG_ADDRESS',
            'straat': 'DEBUG_ADDRESS',
            'postcode': 'DEBUG_POSTAL',
            'postalCode': 'DEBUG_POSTAL',
            'stad': 'DEBUG_CITY',
            'city': 'DEBUG_CITY',
            'btw': 'DEBUG_BTW',
            'companyId': 'DEBUG_BTW',
            'beschrijving': 'DEBUG_DESCRIPTION',
            'date': datetime.now().strftime('%d-%m-%Y'),
        }
        
        print("\n🔍 DEBUG: Scanning all paragraphs for control names...")
        paragraph_count = 0
        found_controls = []
        
        for paragraph in doc.paragraphs:
            paragraph_count += 1
            text = paragraph.text
            
            if text.strip():  # Only check non-empty paragraphs
                print(f"\nParagraph {paragraph_count}: '{text[:100]}{'...' if len(text) > 100 else ''}'")
                
                # Check which controls are found in this paragraph
                found_in_paragraph = []
                for control_name in control_mappings.keys():
                    if control_name in text:
                        found_in_paragraph.append(control_name)
                        found_controls.append((paragraph_count, control_name, text))
                
                if found_in_paragraph:
                    print(f"   🎯 FOUND CONTROLS: {', '.join(found_in_paragraph)}")
        
        print(f"\n🔍 DEBUG: Scanning all table cells...")
        table_count = 0
        
        for table in doc.tables:
            table_count += 1
            print(f"\nTable {table_count}:")
            
            row_count = 0
            for row in table.rows:
                row_count += 1
                cell_count = 0
                
                for cell in row.cells:
                    cell_count += 1
                    
                    for paragraph in cell.paragraphs:
                        text = paragraph.text
                        if text.strip():
                            print(f"   Row {row_count}, Cell {cell_count}: '{text}'")
                            
                            # Check which controls are found in this cell
                            found_in_cell = []
                            for control_name in control_mappings.keys():
                                if control_name in text:
                                    found_in_cell.append(control_name)
                                    found_controls.append((f"Table{table_count}-R{row_count}-C{cell_count}", control_name, text))
                            
                            if found_in_cell:
                                print(f"      🎯 FOUND CONTROLS: {', '.join(found_in_cell)}")
        
        print(f"\n" + "="*60)
        print("📊 SUMMARY OF FOUND CONTROLS:")
        print(f"="*60)
        
        if found_controls:
            for location, control, text in found_controls:
                print(f"{location}: '{control}' in '{text[:50]}{'...' if len(text) > 50 else ''}'")
        else:
            print("❌ NO CONTROLS FOUND!")
            print("This means the control names we're looking for don't exist as plain text.")
            print("They might be:")
            print("1. Inside actual content controls (XML structure)")
            print("2. Split across formatting runs")
            print("3. Using different names than expected")
        
        print(f"\n💡 SUGGESTION:")
        print("Let's check if the controls are structured document tags (content controls)")
        print("rather than plain text that can be replaced with simple string replacement.")
        
    except Exception as e:
        print(f"❌ Error: {e}")

if __name__ == "__main__":
    debug_word_template()